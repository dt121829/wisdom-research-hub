"""Read text out of files the user attaches to the research assistant.

Supports the formats a research desk actually circulates: PDF (including the
reports this app generates), Word, plain text, Markdown and CSV. Everything is
extracted locally — an attached document is never uploaded anywhere except to
the configured AI provider as part of the question the user asks about it.
"""

import csv
import io

TEXT_TYPES = ["pdf", "docx", "txt", "md", "csv"]
IMAGE_TYPES = ["png", "jpg", "jpeg", "webp", "gif", "bmp"]
SUPPORTED = TEXT_TYPES + IMAGE_TYPES

MAX_CHARS = 60_000          # roughly 15k tokens; keeps one document affordable
MAX_IMAGE_PX = 1600         # longest edge; beyond this costs tokens for no more detail


def _from_pdf(data: bytes) -> str:
    import pypdfium2

    doc = pypdfium2.PdfDocument(data)
    try:
        pages = []
        for n, page in enumerate(doc, 1):
            text = page.get_textpage().get_text_range().strip()
            if text:
                pages.append(f"[page {n}]\n{text}")
        return "\n\n".join(pages)
    finally:
        doc.close()


def _from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    rows = list(csv.reader(io.StringIO(text)))
    return "\n".join(" | ".join(r) for r in rows[:400])


def extract_text(name: str, data: bytes) -> tuple[str, str]:
    """Return (text, note). `note` explains any truncation or failure."""
    suffix = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    if suffix in IMAGE_TYPES:
        return "", ""      # images carry no text; they are sent to be looked at
    try:
        if suffix == "pdf":
            text = _from_pdf(data)
        elif suffix == "docx":
            text = _from_docx(data)
        elif suffix == "csv":
            text = _from_csv(data)
        elif suffix in ("txt", "md"):
            text = data.decode("utf-8", errors="replace")
        else:
            return "", f"`{name}`: unsupported file type (.{suffix})."
    except Exception as exc:
        return "", f"`{name}`: could not be read ({type(exc).__name__}: {exc})."

    text = text.strip()
    if not text:
        return "", (f"`{name}`: no text found. If this is a scanned PDF it holds "
                    "images rather than text, which this reader cannot extract.")
    if len(text) > MAX_CHARS:
        return text[:MAX_CHARS], (f"`{name}`: only the first {MAX_CHARS:,} characters "
                                  "were read — the document is longer than one request "
                                  "can carry.")
    return text, ""


def is_image(name: str) -> bool:
    suffix = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    return suffix in IMAGE_TYPES


def normalise_image(data: bytes) -> tuple[bytes, str]:
    """Re-encode an uploaded image as a right-sized PNG. Returns (png, note).

    Screenshots come off a modern display far larger than the model needs;
    shrinking the longest edge cuts token cost without losing legibility of
    chart labels or table figures.
    """
    try:
        from PIL import Image

        img = Image.open(io.BytesIO(data))
        original = f"{img.width}×{img.height}"
        longest = max(img.size)

        # Already a sensible size: send the original bytes rather than re-encoding,
        # which can inflate a well-compressed screenshot for no benefit.
        if longest <= MAX_IMAGE_PX and (img.format or "").upper() in ("PNG", "JPEG"):
            return data, ""

        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        note = ""
        if longest > MAX_IMAGE_PX:
            ratio = MAX_IMAGE_PX / longest
            img = img.resize((max(1, int(img.width * ratio)),
                              max(1, int(img.height * ratio))),
                             Image.LANCZOS)
            note = f"resized from {original} to {img.width}×{img.height}"
        buf = io.BytesIO()
        img.save(buf, format="PNG", optimize=True)
        return buf.getvalue(), note
    except Exception as exc:
        # An unreadable image is better skipped than sent as garbage bytes.
        return b"", f"could not be read ({type(exc).__name__})"


def page_images(name: str, data: bytes, max_pages: int = 6,
                scale: float = 1.6) -> list[bytes]:
    """Render a PDF's pages to PNG so the model can look at charts and tables.

    Text extraction alone loses everything that is drawn rather than written —
    a price chart, a scanned table, a screenshot pasted into a deck.
    """
    if not name.lower().endswith(".pdf"):
        return []
    try:
        import pypdfium2
    except Exception:
        return []

    out = []
    doc = pypdfium2.PdfDocument(data)
    try:
        import io as _io

        for page in list(doc)[:max_pages]:
            buf = _io.BytesIO()
            page.render(scale=scale).to_pil().save(buf, format="PNG", optimize=True)
            out.append(buf.getvalue())
    except Exception:
        return out
    finally:
        doc.close()
    return out


def has_graphics(name: str, data: bytes, sample_pages: int = 6) -> bool:
    """True when a PDF contains drawn content (charts, images) worth looking at."""
    if not name.lower().endswith(".pdf"):
        return False
    try:
        import pypdfium2

        doc = pypdfium2.PdfDocument(data)
        try:
            for page in list(doc)[:sample_pages]:
                for obj in page.get_objects():
                    # 3 = image, 2 = path (charts are drawn as paths).
                    if getattr(obj, "type", None) in (2, 3):
                        return True
        finally:
            doc.close()
    except Exception:
        return False
    return False


def as_context_block(documents: list[dict]) -> str:
    """Format extracted documents for the model, kept separate from live sources."""
    if not documents:
        return ""
    parts = ["<attached_documents>",
             "The user attached these files. They are NOT from the firm's selected "
             "sources: when you use them, say so and name the file."]
    for doc in documents:
        parts.append(f"\n--- FILE: {doc['name']} ({doc['chars']:,} characters) ---")
        parts.append(doc["text"])
    parts.append("</attached_documents>")
    return "\n".join(parts)
