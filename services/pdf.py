"""Render generated reports as professional research-style PDFs.

Pipeline: report Markdown → HTML (python-markdown) → PDF (xhtml2pdf), with an
institutional-research layout: cover header band, document metadata line,
justified body, styled tables, and numbered pages. Price / sector charts are
rendered with matplotlib and embedded as images.

Traditional Chinese output needs a CJK font whose glyphs use TrueType outlines,
because reportlab (the PDF engine) cannot embed PostScript/CFF outlines — which is
what Noto CJK uses. The candidate list below is ordered accordingly, each font is
probed for embeddability, and if embedding still fails at render time the report
falls back to the base font rather than failing outright.
"""

import base64
import io
import re
from datetime import datetime
from functools import lru_cache
from pathlib import Path

_FONT_DIR = Path(__file__).resolve().parent.parent / "data" / "fonts"

# Searched in order. A font dropped into data/fonts/ wins, then the Linux
# packages Streamlit Cloud installs from packages.txt, then local Windows fonts.
# Note: the Windows entries are for development only — those fonts are licensed
# for use on Windows and must not be committed to the repository.
def _bundled_fonts() -> list[Path]:
    """Any font dropped into data/fonts/ — checked before the system ones."""
    if not _FONT_DIR.exists():
        return []
    return sorted(_FONT_DIR.glob("*.ttf")) + sorted(_FONT_DIR.glob("*.otf"))


# Ordered best-first. reportlab can only embed TrueType-outline fonts, so the
# TrueType CJK fonts (WenQuanYi, AR PL, Windows kaiu) come before the Noto CJK
# ttc, whose PostScript/CFF outlines reportlab rejects. Each candidate is still
# validated below, so an un-embeddable one is skipped rather than crashing.
_CJK_FONT_CANDIDATES = [
    *_bundled_fonts(),
    # Debian/Ubuntu (Streamlit Community Cloud) — TrueType outlines, embeddable.
    Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
    Path("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"),
    Path("/usr/share/fonts/truetype/arphic/uming.ttc"),
    Path("/usr/share/fonts/truetype/arphic/ukai.ttc"),
    # Local Windows development — also TrueType.
    Path(r"C:\Windows\Fonts\kaiu.ttf"),      # DFKai-SB — Traditional Chinese
    Path(r"C:\Windows\Fonts\msjh.ttf"),
    Path(r"C:\Windows\Fonts\mingliu.ttf"),
    # Last resort: Noto CJK. Usually rejected by reportlab (CFF outlines) and so
    # skipped by the validation, but kept in case a TrueType build is present.
    Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
    Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
]


def _reportlab_can_embed(path: Path) -> bool:
    """True only if reportlab can actually embed this font into a PDF.

    reportlab supports TrueType (glyf) outlines but not PostScript/CFF, and the
    only reliable test is to try loading it — a bare exists() check let the
    un-embeddable Noto CJK ttc through and crashed the whole render.
    """
    try:
        from reportlab.pdfbase.ttfonts import TTFont

        TTFont("probe", str(path))
        return True
    except Exception:
        return False


@lru_cache(maxsize=1)
def _cjk_font_path() -> str | None:
    for p in _CJK_FONT_CANDIDATES:
        if p.exists() and _reportlab_can_embed(p):
            # Forward slashes: backslash paths get misread as URLs by the CSS parser.
            return str(p).replace("\\", "/")
    return None


# ------------------------------------------------------------------- charts

def _chart_style(ax, fig):
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    for spine in ("left", "bottom"):
        ax.spines[spine].set_color("#c3c2b7")
    ax.tick_params(colors="#898781", labelsize=8)
    ax.grid(True, axis="y", color="#e1e0d9", linewidth=0.6)
    ax.set_axisbelow(True)


def price_chart_png(dates, values, title: str) -> bytes | None:
    """Line chart of closes, styled to match the app. Returns PNG bytes."""
    if not values or len(values) < 2:
        return None
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(6.4, 2.6), dpi=150)
    x = range(len(values))
    ax.plot(x, values, color="#2a78d6", linewidth=1.6)
    ax.fill_between(x, values, min(values), color="#2a78d6", alpha=0.07)
    # ~6 date ticks across the range
    step = max(1, len(dates) // 6)
    ticks = list(range(0, len(dates), step))
    ax.set_xticks(ticks)
    ax.set_xticklabels([dates[i] for i in ticks], fontsize=7)
    ax.set_title(title, fontsize=9, color="#0b0b0b", loc="left", pad=6)
    _chart_style(ax, fig)
    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def sector_chart_png(sectors, title: str, highlight: str = "") -> bytes | None:
    """Horizontal bar chart of sector returns. Returns PNG bytes.

    `highlight` names the sector the report is about: it is drawn in the accent
    colour with a bold label so it stands out from the rest of the set.
    """
    if not sectors:
        return None
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    s_sorted = sorted(sectors, key=lambda s: s["change"])
    names = [s["name"] for s in s_sorted]
    changes = [s["change"] for s in s_sorted]
    hl = (highlight or "").strip().lower()

    def is_target(name):
        return bool(hl) and (hl in name.lower() or name.lower() in hl)

    colors = ["#2a78d6" if is_target(n)
              else ("#0ca30c" if c >= 0 else "#d03b3b")
              for n, c in zip(names, changes)]

    fig, ax = plt.subplots(figsize=(6.4, 3.2), dpi=150)
    bars = ax.barh(names, changes, color=colors, height=0.62)
    ax.set_title(title, fontsize=9, color="#0b0b0b", loc="left", pad=6)
    ax.xaxis.set_major_formatter(lambda v, _: f"{v:+.0f}%")
    _chart_style(ax, fig)
    ax.grid(True, axis="x", color="#e1e0d9", linewidth=0.6)
    ax.grid(False, axis="y")

    for label, name, bar, change in zip(ax.get_yticklabels(), names, bars, changes):
        if is_target(name):
            label.set_fontweight("bold")
            label.set_color("#10365c")
            bar.set_edgecolor("#10365c")
            bar.set_linewidth(0.9)
            ax.annotate(f"  {change:+.1f}%  ← this report",
                        xy=(change, bar.get_y() + bar.get_height() / 2),
                        va="center", ha="left" if change >= 0 else "right",
                        fontsize=7.5, color="#10365c", fontweight="bold")
    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def peer_chart_png(rows, subject: str, title: str) -> bytes | None:
    """Return-comparison bars for a stock against its peers. Returns PNG bytes."""
    if not rows or len(rows) < 2:
        return None
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    rows = sorted(rows, key=lambda r: r["change"])
    names = [r["name"] for r in rows]
    changes = [r["change"] for r in rows]
    colors = ["#2a78d6" if n.upper() == subject.upper()
              else ("#0ca30c" if c >= 0 else "#d03b3b")
              for n, c in zip(names, changes)]

    fig, ax = plt.subplots(figsize=(6.4, 2.8), dpi=150)
    bars = ax.barh(names, changes, color=colors, height=0.6)
    ax.set_title(title, fontsize=9, color="#0b0b0b", loc="left", pad=6)
    ax.xaxis.set_major_formatter(lambda v, _: f"{v:+.0f}%")
    _chart_style(ax, fig)
    ax.grid(True, axis="x", color="#e1e0d9", linewidth=0.6)
    ax.grid(False, axis="y")

    avg = sum(changes) / len(changes)
    ax.axvline(avg, linestyle="--", color="#898781", linewidth=0.9)
    ax.annotate(f"peer avg {avg:+.1f}%", xy=(avg, len(rows) - 0.4),
                fontsize=7, color="#52514e", ha="center")
    for label, name, bar in zip(ax.get_yticklabels(), names, bars):
        if name.upper() == subject.upper():
            label.set_fontweight("bold")
            label.set_color("#10365c")
            bar.set_edgecolor("#10365c")
            bar.set_linewidth(0.9)
    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


# ---------------------------------------------------------------------- html

def _clean_md(md_text: str) -> str:
    """Undo Streamlit-specific escaping and swap glyphs base PDF fonts lack."""
    replacements = {
        "\\$": "$",
        "\u2011": "-",   # non-breaking hyphen
        "\u2212": "-",   # minus sign
        "\u00a0": " ",   # no-break space
        "\u202f": " ",   # narrow no-break space
        "\ufe0f": "",    # emoji variation selector
    }
    # The PDF base fonts carry no emoji, so these would print as hollow boxes.
    for emoji, sub in (("\u26a0", "!"), ("\u2694", ""), ("\U0001f4ca", ""),
                       ("\U0001f4c5", ""), ("\U0001f4ce", ""), ("\U0001f3e6", ""),
                       ("\U0001f3db", ""), ("\U0001f7e2", ""), ("\U0001f534", ""),
                       ("\u26aa", ""), ("\U0001f7e1", "")):
        replacements[emoji] = sub
    for old, new_ in replacements.items():
        md_text = md_text.replace(old, new_)
    # Tidy up spacing left behind by the removals.
    md_text = re.sub(r"[ \t]{2,}", " ", md_text)
    return re.sub(r"(^|\n)([-*] )\s+", r"\1\2", md_text)


def _img_tag(png: bytes, caption: str = "") -> str:
    b64 = base64.b64encode(png).decode("ascii")
    cap = f'<p class="figcaption">{caption}</p>' if caption else ""
    return f'<div class="figure"><img src="data:image/png;base64,{b64}"/>{cap}</div>'


_CSS = """
@page {{
    size: A4;
    margin: 2.2cm 1.9cm 2.4cm 1.9cm;
    @frame footer_frame {{
        -pdf-frame-content: footer_content;
        left: 1.9cm; width: 17.2cm; top: 27.6cm; height: 0.9cm;
    }}
}}
{font_face}
body {{ font-family: {body_font}; font-size: 9.5pt; color: #1a1a18; line-height: 1.55; }}
.band {{ background-color: #10365c; color: #ffffff; padding: 10px 14px 8px 14px; }}
.band .firm {{ font-size: 13pt; font-weight: bold; color: #ffffff; }}
.band .dept {{ font-size: 8pt; color: #b8cbe2; }}
.meta {{ font-size: 8pt; color: #52514e; border-bottom: 1.5pt solid #10365c;
         padding: 5px 0 6px 0; margin-bottom: 14px; }}
h1 {{ font-size: 17pt; color: #10365c; margin: 12px 0 6px 0; line-height: 1.25; }}
h2 {{ font-size: 11.5pt; color: #10365c; border-bottom: 0.75pt solid #c3c2b7;
      padding-bottom: 3px; margin: 18px 0 7px 0; -pdf-keep-with-next: true; }}
h3 {{ font-size: 10pt; color: #1a1a18; margin: 13px 0 4px 0; -pdf-keep-with-next: true; }}
p {{ margin: 5px 0; word-wrap: break-word; }}
a {{ color: #2a78d6; text-decoration: none; }}
ol li, ul li {{ word-wrap: break-word; }}
li {{ margin: 4px 0; word-wrap: break-word; }}
strong {{ color: #10365c; }}
em {{ color: #52514e; }}
/* A standalone italic line is a disclaimer/footnote, not body copy. */
p > em:only-child {{ font-size: 7.5pt; color: #898781; }}
/* Tables: fixed layout keeps long cell text inside its column instead of
   pushing the table past the page margin. Left-aligned, never justified. */
table {{ width: 100%; border-collapse: collapse; margin: 9px 0; font-size: 8pt;
         -pdf-keep-in-frame-mode: shrink; }}
th {{ background-color: #10365c; color: #ffffff; padding: 5px 6px; text-align: left;
      font-size: 8pt; }}
td {{ padding: 5px 6px; border-bottom: 0.5pt solid #e1e0d9; vertical-align: top;
      text-align: left; word-wrap: break-word; }}
blockquote {{ color: #52514e; border-left: 2pt solid #c3c2b7; padding-left: 8px;
              margin: 6px 0 6px 4px; }}
/* Charts sit at ~2/3 width so a short report keeps its page budget for text. */
.figure {{ margin: 8px 0; }}
.figure img {{ width: 11.5cm; }}
.figcaption {{ font-size: 7.5pt; color: #898781; margin: 2px 0 0 0; }}
.disclaimer {{ font-size: 7.5pt; color: #898781; border-top: 0.5pt solid #c3c2b7;
               padding-top: 6px; margin-top: 18px; }}
#footer_content {{ font-size: 7.5pt; color: #898781; text-align: center; }}
"""


def markdown_to_pdf(md_text: str, *, subtitle: str = "", charts_html: str = "",
                    firm: str = "Wisdom Family Office") -> bytes:
    """Convert report markdown to a professionally formatted PDF. Returns bytes."""
    import markdown as md_lib
    from xhtml2pdf import pisa

    # Windows workaround: xhtml2pdf copies url('...') resources to a temp file it
    # still holds open, which reportlab then cannot read (sharing violation). Our
    # font src is a local path, so hand the path straight through instead.
    try:
        from xhtml2pdf.files import pisaFileObject
        pisaFileObject.getNamedFile = lambda self: (self.uri or "").replace("file:///", "")
    except Exception:
        pass

    md_text = _clean_md(md_text)
    body = md_lib.markdown(md_text, extensions=["tables", "fenced_code", "sane_lists"])

    # Insert the charts right after the first <h1> (or at the top when there is none).
    if charts_html:
        if "</h1>" in body:
            body = body.replace("</h1>", "</h1>" + charts_html, 1)
        else:
            body = charts_html + body

    cjk = _cjk_font_path()
    has_cjk_text = bool(re.search(r"[\u4e00-\u9fff]", md_text))
    generated = datetime.now().strftime("%B %d, %Y · %H:%M")

    def _render(font_face: str, body_font: str) -> bytes:
        css = _CSS.format(font_face=font_face, body_font=body_font)
        html = f"""<html><head><style>{css}</style></head><body>
<div class="band">
  <span class="firm">{firm}</span><br/>
  <span class="dept">Investment Research · AI-Synthesised Report</span>
</div>
<div class="meta">{subtitle or 'Research report'} &nbsp;|&nbsp; Generated {generated}
 &nbsp;|&nbsp; Sources: Seeking Alpha · Yahoo Finance · CNBC · SumZero</div>
{body}
<p class="disclaimer">This document is an AI-generated synthesis of third-party research
from the firm's selected sources, prepared for internal use by {firm}. It is not
investment advice and should not be relied upon as the sole basis for any investment
decision. Figures are as of the generation time above.</p>
<div id="footer_content">{firm} · Confidential — internal use only · Page <pdf:pagenumber/>
 of <pdf:pagecount/></div>
</body></html>"""
        out = io.BytesIO()
        result = pisa.CreatePDF(io.StringIO(html), dest=out, encoding="utf-8")
        if result.err:
            raise RuntimeError("PDF rendering failed")
        return out.getvalue()

    # Try the CJK font first. The up-front probe cannot fully guarantee that
    # xhtml2pdf will embed a .ttc collection at render time, so if that fails we
    # fall back to the base font — a document with Chinese shown as boxes still
    # beats no document at all.
    if cjk and has_cjk_text:
        try:
            return _render(f"@font-face {{ font-family: CJK; src: url('{cjk}'); }}",
                           "CJK")
        except Exception:
            pass
    return _render("", "Helvetica")


def page_images(pdf_bytes: bytes, scale: float = 2.0) -> list[bytes]:
    """Rasterise each PDF page to PNG bytes.

    The app displays these directly, so the in-app view is the real rendered
    document rather than a browser plugin that may be blocked or blank.
    """
    import pypdfium2

    doc = pypdfium2.PdfDocument(pdf_bytes)
    out = []
    try:
        for page in doc:
            buf = io.BytesIO()
            page.render(scale=scale).to_pil().save(buf, format="PNG")
            out.append(buf.getvalue())
    finally:
        doc.close()
    return out


# ---------------------------------------------------------------------- docx

def markdown_to_docx(md_text: str, *, subtitle: str = "",
                     firm: str = "Wisdom Family Office") -> bytes:
    """Convert report markdown to a Word document. Returns .docx bytes."""
    import markdown as md_lib
    from docx import Document
    from docx.shared import Pt, RGBColor
    from htmldocx import HtmlToDocx

    html = md_lib.markdown(_clean_md(md_text),
                           extensions=["tables", "fenced_code", "sane_lists"])
    # htmldocx chokes on raw anchors; strip them (Word keeps the Sources list).
    html = re.sub(r"<a name=[^>]*></a>", "", html)

    doc = Document()
    head = doc.add_paragraph()
    run = head.add_run(firm)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x10, 0x36, 0x5C)
    sub = doc.add_paragraph()
    run = sub.add_run(
        (subtitle or "Research report")
        + f" | Generated {datetime.now():%B %d, %Y · %H:%M}"
        + " | Sources: Seeking Alpha · Yahoo Finance · CNBC · SumZero")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x52, 0x51, 0x4E)

    HtmlToDocx().add_html_to_document(html, doc)

    tail = doc.add_paragraph()
    run = tail.add_run(
        f"This document is an AI-generated synthesis of third-party research from the "
        f"firm's selected sources, prepared for internal use by {firm}. It is not "
        "investment advice.")
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x89, 0x87, 0x81)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
